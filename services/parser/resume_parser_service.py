# services/parser/resume_parser_service.py

import os
import uuid
import tempfile
import docx
import re
import json
import logging
from django.conf import settings
import openai
import google.generativeai as genai
from django.utils import timezone
from datetime import datetime
from urllib.parse import urlparse
import traceback

# Setup logging
logger = logging.getLogger(__name__)

# Try to import PyMuPDF (fitz) for enhanced PDF processing
try:
    import fitz  # PyMuPDF
except ImportError:
    logger.warning("PyMuPDF (fitz) not installed. PDF link extraction disabled. Install with: pip install PyMuPDF")
    fitz = None


class ResumeParserService:
    """
    Service for parsing resumes and extracting structured data from them.
    Uses user-provided API keys instead of global settings.
    """

    def __init__(self, user_openai_key=None, user_gemini_key=None):
        """
        Initialize the parser service with user-specific API keys.
        Args:
            user_openai_key (str, optional): User's OpenAI API key
            user_gemini_key (str, optional): User's Google Gemini API key
        """
        self.user_openai_key = user_openai_key
        self.user_gemini_key = user_gemini_key
        self.openai_client = None
        self.gemini_model = None

        # Initialize clients if keys are provided
        if self.user_openai_key:
            try:
                self.openai_client = openai.OpenAI(api_key=self.user_openai_key)
                logger.info("OpenAI client initialized with user API key")
            except Exception as e:
                logger.error(f"Failed to initialize OpenAI client: {str(e)}")

        if self.user_gemini_key:
            try:
                # Configure Gemini with user API key
                genai.configure(api_key=self.user_gemini_key)
                model_name = getattr(settings, 'GEMINI_MODEL_NAME', 'gemini-1.5-flash')
                self.gemini_model = genai.GenerativeModel(model_name)
                logger.info(f"Gemini model '{model_name}' initialized with user API key")
            except Exception as e:
                logger.error(f"Failed to initialize Gemini model: {str(e)}")

    def parse_resume(self, file_path, file_type=None, ai_parsing_enabled=True, ai_provider="gemini"):
        """
        Parse a resume file and extract structured data.

        Args:
            file_path (str): Path to the resume file
            file_type (str, optional): File extension to override auto-detection
            ai_parsing_enabled (bool): Whether to use AI for parsing
            ai_provider (str): Which AI provider to use ("gemini" or "chatgpt")

        Returns:
            dict: Parsed resume data in structured format
        """
        try:
            # Determine file type if not provided
            if not file_type:
                file_type = os.path.splitext(file_path)[1].lower().strip('.')

            # Extract text and links from the resume file
            resume_text, extracted_links = self._extract_text_and_links(file_path, file_type)

            # Check if text extraction was successful
            if not resume_text or len(resume_text.strip()) < 50:
                logger.warning(f"Extracted text is too short ({len(resume_text)} chars). Extraction might have failed.")
                return {"error": "Could not extract sufficient text from the resume. Please check the file format."}

            # Parse the resume text
            if ai_parsing_enabled:
                if ai_provider == "chatgpt" and self.openai_client:
                    parsed_data = self._parse_with_openai(resume_text, extracted_links)
                elif ai_provider == "gemini" and self.gemini_model:
                    parsed_data = self._parse_with_gemini(resume_text, extracted_links)
                else:
                    logger.warning(f"AI provider '{ai_provider}' not available or no API key. Using basic parsing.")
                    parsed_data = self._basic_resume_parsing(resume_text)
            else:
                logger.info("AI parsing disabled. Using basic parsing.")
                parsed_data = self._basic_resume_parsing(resume_text)

            return parsed_data

        except Exception as e:
            logger.error(f"Error parsing resume: {str(e)}")
            traceback.print_exc()
            return {"error": f"Resume parsing failed: {str(e)}"}

    def _extract_text_and_links(self, file_path, file_type):
        """
        Extract text and links from a resume file based on file type.

        Args:
            file_path (str): Path to the file
            file_type (str): File extension (pdf, docx, doc, txt)

        Returns:
            tuple: (extracted_text, extracted_links)
        """
        if file_type == 'pdf':
            return self._extract_from_pdf(file_path)
        elif file_type in ['docx', 'doc']:
            return self._extract_from_docx(file_path), []
        elif file_type in ['txt', 'text', 'odt']:
            return self._extract_from_txt(file_path), []
        else:
            logger.warning(f"Unsupported file type: {file_type}")
            return "", []

    def _extract_from_pdf(self, file_path):
        """Extract text and links from a PDF file."""
        text_content = ""
        extracted_links = []

        # Use PyMuPDF if available
        if fitz:
            doc = None
            try:
                doc = fitz.open(file_path)
                for page_num in range(len(doc)):
                    page = doc.load_page(page_num)
                    text_content += page.get_text("text") + "\n"

                    # Extract links from the page
                    links = page.get_links()
                    for link in links:
                        if link.get('kind') == fitz.LINK_URI and 'uri' in link:
                            uri = link['uri']
                            if uri and uri not in extracted_links:
                                if isinstance(uri, str) and (uri.startswith('http') or uri.startswith('mailto:')):
                                    extracted_links.append(uri)
                                elif isinstance(uri, str) and '@' in uri and '.' in uri and not uri.startswith(
                                        'mailto:'):
                                    if re.match(r"[^@]+@[^@]+\.[^@]+", uri):
                                        extracted_links.append(f"mailto:{uri}")
            except Exception as e:
                logger.error(f"Error extracting text/links from PDF with PyMuPDF: {str(e)}")
                return "", []
            finally:
                if doc:
                    try:
                        doc.close()
                    except Exception as e:
                        logger.error(f"Error closing PDF document: {str(e)}")

            logger.info(f"PyMuPDF extracted text length: {len(text_content)}, links: {len(extracted_links)}")
            return text_content, extracted_links

        # Fallback to PyPDF2 if PyMuPDF is not available
        else:
            try:
                import PyPDF2
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)

                    # Check if PDF is encrypted
                    if pdf_reader.is_encrypted:
                        try:
                            if pdf_reader.decrypt('') == 0:  # Failed decryption
                                logger.warning(f"Could not decrypt PDF {file_path}")
                                return "", []
                        except Exception as e:
                            logger.error(f"Error decrypting PDF {file_path}: {str(e)}")
                            return "", []

                    # Extract text from each page
                    for page in pdf_reader.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text_content += page_text + "\n"

                logger.info(f"PyPDF2 extracted text length: {len(text_content)}")
                return text_content, []  # PyPDF2 doesn't extract links
            except Exception as e:
                logger.error(f"Error during fallback PDF text extraction with PyPDF2: {str(e)}")
                return "", []

    def _extract_from_docx(self, file_path):
        """Extract text from a DOCX file."""
        text = ""
        try:
            doc = docx.Document(file_path)
            for para in doc.paragraphs:
                text += para.text + "\n"
        except Exception as e:
            logger.error(f"Error extracting text from DOCX: {str(e)}")
        return text

    def _extract_from_txt(self, file_path):
        """Extract text from a plain text file."""
        text = ""
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
                text = file.read()
        except Exception as e:
            logger.error(f"Error reading TXT file {file_path}: {str(e)}")
            try:
                # Fallback to latin-1 encoding
                with open(file_path, 'r', encoding='latin-1', errors='ignore') as file:
                    text = file.read()
            except Exception as e:
                logger.error(f"Error reading TXT with latin-1: {str(e)}")
        return text

    def _parse_with_openai(self, resume_text, extracted_links):
        """Parse resume text using OpenAI API."""
        context_text = self._prepare_context_text(resume_text, extracted_links)
        prompt = self._get_parsing_prompt(context_text)

        try:
            response = self.openai_client.chat.completions.create(
                model=getattr(settings, 'OPENAI_MODEL', 'gpt-4o-mini'),
                messages=[
                    {"role": "system",
                     "content": "Parse resume text into the specified JSON format. Ensure all values are correct types. List individual skills, do not categorize them. Output only valid JSON."},
                    {"role": "user", "content": prompt}
                ],
                temperature=0.1,
                max_tokens=4000,
                response_format={"type": "json_object"}
            )

            result = response.choices[0].message.content

            # Clean the response
            result = re.sub(r'^```json\s*|\s*```$', '', result.strip(), flags=re.MULTILINE)
            result = re.sub(r'^\s*//.*$', '', result, flags=re.MULTILINE)

            # Parse the JSON
            parsed_data = json.loads(result)

            # Validate the structure
            if not isinstance(parsed_data, dict):
                logger.error(f"Parsed data is not a dictionary (type: {type(parsed_data)})")
                return self._basic_resume_parsing(resume_text)

            return parsed_data

        except json.JSONDecodeError as e:
            logger.error(f"JSON decode error from OpenAI: {str(e)}")
            return self._basic_resume_parsing(resume_text)
        except Exception as e:
            logger.error(f"Error parsing with OpenAI: {str(e)}")
            return self._basic_resume_parsing(resume_text)

    def _parse_with_gemini(self, resume_text, extracted_links):
        """Parse resume text using Google Gemini API."""
        context_text = self._prepare_context_text(resume_text, extracted_links)
        prompt = self._get_parsing_prompt(context_text)

        try:
            response = self.gemini_model.generate_content(
                prompt,
                generation_config=genai.types.GenerationConfig(
                    temperature=0.2,
                    response_mime_type="application/json"
                )
            )

            result = response.text

            # Clean the response
            result = re.sub(r'^```json\s*|\s*```$', '', result.strip(), flags=re.MULTILINE)
            result = re.sub(r'^\s*//.*$', '', result, flags=re.MULTILINE)

            # Parse the JSON
            parsed_data = json.loads(result)

            # Validate the structure
            if not isinstance(parsed_data, dict):
                logger.error(f"Parsed data is not a dictionary (type: {type(parsed_data)})")
                return self._basic_resume_parsing(resume_text)

            return parsed_data

        except json.JSONDecodeError as e:
            logger.error(f"JSON decode error from Gemini: {str(e)}")
            return self._basic_resume_parsing(resume_text)
        except Exception as e:
            logger.error(f"Error parsing with Gemini: {str(e)}")
            return self._basic_resume_parsing(resume_text)

    def _prepare_context_text(self, resume_text, extracted_links):
        """Prepare the context text with resume text and extracted links."""
        context_text = resume_text

        if extracted_links:
            context_text += "\n\n--- Extracted Links (Prioritize these for URLs/Email) ---\n"
            for link in extracted_links:
                context_text += f"- {link}\n"
            context_text += "--- End Extracted Links ---\n"

        return context_text

    def _get_parsing_prompt(self, context_text):
        """
        Returns the prompt for the AI model to parse resume data.
        """
        return f"""
        You are an expert resume parser. Your task is to extract all relevant information from the provided resume text
        and structure it into a comprehensive JSON object. Ensure all fields are extracted accurately and completely.

        For dates, use 'YYYY-MM-DD' format if specific day is available, otherwise 'YYYY-MM' or just 'YYYY'. If a date range is given, use the earliest date for 'Start date' and the latest date for 'End date' or 'Completion date'. If a position is current, set 'Is current job' to true and 'End date' to null.

        The JSON output should strictly follow the following structure:
        ```json
        {{
          "Personal Information": {{
            "First name": "string",
            "Middle name": "string | null",
            "Last name": "string",
            "Email": "string (use mailto: links or email addresses found)",
            "Phone number": "string",
            "Address": "string (e.g., 'City, State/Province, Country') | null",
            "LinkedIn URL": "string (full URL) | null",
            "GitHub URL": "string (full URL) | null",
            "Portfolio URL": "string (full URL) | null"
          }},
          "Professional Summary": "string | null",
          "Skills": [
            {{
              "Skill name": "string",
              "Category": "string (e.g., 'Programming Languages', 'Frameworks', 'Tools', 'Soft Skills', 'Other') | null",
              "Other category": "string | null (if 'Category' is 'Other', specify here)",
              "Estimated proficiency level": "integer (0-100, where 100 is expert and 0 is basic understanding)"
            }}
          ] | null,
          "Work Experience": [
            {{
              "Job title": "string",
              "Employer/Company name": "string",
              "Location": "string | null",
              "Start date": "string | null (YYYY-MM-DD or YYYY-MM or YYYY)",
              "End date": "string | null (YYYY-MM-DD or YYYY-MM or YYYY)",
              "Is current job": "boolean",
              "Bullet points": ["string", ...] | null
            }}
          ] | null,
          "Education": [
            {{
              "School name": "string",
              "Location": "string | null",
              "Degree": "string",
              "Degree type": "string (e.g., 'High School Diploma', 'Associate', 'Bachelor', 'Master', 'Doctorate', 'Certificate', 'Other')",
              "Field of study": "string | null",
              "Graduation date": "string | null (YYYY-MM-DD or YYYY-MM or YYYY)",
              "GPA": "float | null",
              "Description": "string | null"
            }}
          ] | null,
          "Projects": [
            {{
              "Project name": "string",
              "Summary/description": "string | null",
              "Start date": "string | null (YYYY-MM-DD or YYYY-MM or YYYY)",
              "Completion date": "string | null (YYYY-MM-DD or YYYY-MM or YYYY)",
              "Project URL": "string | null",
              "GitHub URL": "string | null",
              "Bullet points": ["string", ...] | null
            }}
          ] | null,
          "Certifications": [
            {{
              "Name": "string",
              "Institute/Issuing organization": "string | null",
              "Completion date": "string | null (YYYY-MM-DD or YYYY-MM or YYYY)",
              "Expiration date": "string | null (YYYY-MM-DD or YYYY-MM or YYYY)",
              "Score": "string | null",
              "URL/Link": "string | null",
              "Description": "string | null"
            }}
          ] | null,
          "Languages": [
            {{
              "Language name": "string",
              "Proficiency (0-100 where 100 is native and 0 is basic)": "integer | null"
            }}
          ] | null,
          "Custom Sections": [
            {{
              "Section title": "string",
              "Entries": [
                {{
                  "Entry title": "string",
                  "Description": "string | null",
                  "Start date": "string | null (YYYY-MM-DD or YYYY-MM or YYYY)",
                  "End date": "string | null (YYYY-MM-DD or YYYY-MM or YYYY)",
                  "Link": "string | null"
                }}
              ] | null
            }}
          ] | null
        }}
        ```
        - If a section is not present in the resume, its corresponding key in the JSON should be `null`.
        - For lists (Skills, Work Experience, etc.), if no entries are found, the list should be `null`.
        - Ensure all URLs are full and valid.
        - Pay close attention to date formats.
        - Bullet points should be extracted verbatim as individual strings in a list.
        - For skills, list individual skills. Do not group skills into broad categories unless specifically asked in the "Category" field.
        """

    def _basic_resume_parsing(self, resume_text):
        """Basic fallback parsing for when AI parsing fails."""
        if not isinstance(resume_text, str):
            resume_text = str(resume_text)

        email = None
        phone = None
        linkedin = None
        github = None
        portfolio = None
        first_name = ""
        last_name = ""
        middle_name = None

        # Extract email
        try:
            email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
            emails = re.findall(email_pattern, resume_text)
            email = emails[0] if emails else None
        except Exception as e:
            logger.error(f"Error parsing email: {str(e)}")

        # Extract phone
        try:
            phone_pattern = r'\b(?:\+\d{1,3}[-\.\s]?)?\(?\d{3}\)?[-\.\s]?\d{3}[-\.\s]?\d{4}\b'
            phones = re.findall(phone_pattern, resume_text)
            phone = phones[0] if phones else None
        except Exception as e:
            logger.error(f"Error parsing phone: {str(e)}")

        # Extract URLs
        try:
            url_pattern = r'https?://(?:www\.)?[-a-zA-Z0-9@:%._\+~#=]{1,256}\.[a-zA-Z0-9()]{1,6}\b(?:[-a-zA-Z0-9()@:%_\+.~#?&//=]*)'
            urls = re.findall(url_pattern, resume_text)

            for url in urls:
                try:
                    parsed = urlparse(url)
                    scheme = parsed.scheme or 'https'
                    netloc = parsed.netloc.replace('www.', '')
                    path = parsed.path or ''
                    clean_url = f"{scheme}://{netloc}{path}".rstrip('/')

                    if 'linkedin.com/in/' in clean_url and not linkedin:
                        linkedin = clean_url
                    elif 'github.com/' in clean_url and not github:
                        github = clean_url
                    elif not any(x in clean_url for x in ['linkedin.com', 'github.com']) and not portfolio:
                        if not any(ext in clean_url for ext in ['.pdf', '.png', '.jpg', '.jpeg', '.gif']):
                            portfolio = clean_url
                except ValueError:
                    logger.warning(f"Skipping malformed URL: {url}")
                    continue
        except Exception as e:
            logger.error(f"Error parsing URLs: {str(e)}")

        # Extract name
        try:
            lines = resume_text.strip().split('\n')
            name = lines[0].strip() if lines else ""
            name_parts = name.split()
            first_name = name_parts[0] if name_parts else ""
            last_name = name_parts[-1] if len(name_parts) > 1 else ""
            middle_name = " ".join(name_parts[1:-1]) if len(name_parts) > 2 else None
        except Exception as e:
            logger.error(f"Error parsing name: {str(e)}")

        # Return the default dictionary structure
        return {
            "Personal Information": {
                "First name": first_name or None,
                "Middle name": middle_name,
                "Last name": last_name or None,
                "Email": email,
                "Phone number": phone,
                "Address": None,
                "LinkedIn URL": linkedin,
                "GitHub URL": github,
                "Portfolio URL": portfolio
            },
            "Professional Summary": None,
            "Skills": [],
            "Work Experience": [],
            "Education": [],
            "Projects": [],
            "Certifications": [],
            "Languages": [],
            "Custom Sections": []
        }

    @staticmethod
    def format_date(date_str):
        """Convert various date string formats to YYYY-MM-DD."""
        if not date_str or not isinstance(date_str, str) or str(date_str).lower() == 'present':
            return None

        date_str = str(date_str).strip()

        # Try common date formats
        formats = [
            '%Y-%m-%d', '%m/%d/%Y', '%m-%d-%Y', '%d-%b-%Y', '%d-%B-%Y',
            '%Y/%m/%d', '%d %b %Y', '%d %B %Y', '%b %d, %Y', '%B %d, %Y',
            '%Y-%m', '%m/%Y', '%b %Y', '%B %Y', '%Y',
        ]

        for fmt in formats:
            try:
                dt = datetime.strptime(date_str, fmt)

                # Handle formats with varying specificity
                if fmt in ['%Y-%m', '%m/%Y', '%b %Y', '%B %Y']:
                    return dt.strftime('%Y-%m-01')  # Default to the 1st of the month
                elif fmt == '%Y':
                    return dt.strftime('%Y-01-01')  # Default to Jan 1st of the year
                else:
                    return dt.strftime('%Y-%m-%d')  # Full date
            except ValueError:
                continue

        # Fallback attempts
        try:
            # Try to match "Month Year" format
            match_month_year = re.match(r'([A-Za-z]{3,})\s+(\d{4})', date_str)
            if match_month_year:
                month_str, year_str = match_month_year.groups()

                try:
                    dt = datetime.strptime(f"{month_str} 1 {year_str}", '%b 1 %Y')
                    return dt.strftime('%Y-%m-01')
                except ValueError:
                    try:
                        dt = datetime.strptime(f"{month_str} 1 {year_str}", '%B 1 %Y')
                        return dt.strftime('%Y-%m-01')
                    except ValueError:
                        pass

            # Try to match just a year
            match_year = re.match(r'(\d{4})', date_str)
            if match_year:
                year = match_year.group(1)
                return f"{year}-01-01"
        except Exception as e:
            logger.error(f"Error during fallback date parsing for '{date_str}': {str(e)}")

        logger.warning(f"Could not parse date string: {date_str}")
        return None

    @staticmethod
    def format_url(url_str):
        """Adds https:// if missing and cleans the URL."""
        if not url_str or not isinstance(url_str, str):
            return None

        url_str = url_str.strip()
        if not url_str:
            return None

        try:
            # Add scheme if missing, default to https
            if not re.match(r"^(?:[a-z]+:)?//", url_str, re.IGNORECASE) and '.' in url_str:
                url_str = 'https://' + url_str

            parsed = urlparse(url_str)
            scheme = parsed.scheme
            netloc = parsed.netloc
            path = parsed.path or ''

            # Basic validation
            if not scheme or not netloc or '.' not in netloc:
                # Check if it's an email
                if '@' in url_str and '.' in url_str and not url_str.startswith('mailto:'):
                    if re.match(r"[^@]+@[^@]+\.[^@]+", url_str):
                        return f"mailto:{url_str}"
                return None

            # Reconstruct URL, removing trailing slash if path is not just '/'
            clean_url = f"{scheme}://{netloc}{path}"
            if clean_url.endswith('/') and len(path) > 1:
                clean_url = clean_url.rstrip('/')

            return clean_url
        except Exception as e:
            logger.error(f"Error formatting URL '{url_str}': {str(e)}")
            return None

    @staticmethod
    def format_location(location_data):
        """Formats location data into a string."""
        if isinstance(location_data, dict):
            city = location_data.get('City', '')
            state = location_data.get('State/Province', '')
            country = location_data.get('Country', '')

            parts = [str(part).strip() for part in [city, state, country]
                     if part and isinstance(part, str) and str(part).strip()]
            return ", ".join(parts) if parts else ""

        elif isinstance(location_data, str):
            return location_data.strip()

        elif isinstance(location_data, (list, tuple)):
            parts = [str(item).strip() for item in location_data
                     if item and isinstance(item, str) and str(item).strip()]
            return ", ".join(parts) if parts else ""

        else:
            return ""

    @staticmethod
    def safe_strip(value, default=''):
        """Safely strips a value if it's a string, joins if list/tuple, returns default otherwise."""
        if isinstance(value, str):
            return value.strip()

        elif isinstance(value, (list, tuple)):
            string_items = [str(item).strip() for item in value
                            if isinstance(item, str) and str(item).strip()]
            return " ".join(string_items) if string_items else default

        elif value is None:
            return default

        else:
            try:
                return str(value).strip()
            except:
                return default

# # services/parser/resume_parser_service.py
#
# import os
# import uuid
# import tempfile
# import docx
# import re
# import json
# import logging
# from django.conf import settings
# import openai
# import google.generativeai as genai
# from django.utils import timezone
# from datetime import datetime
# from urllib.parse import urlparse
# import traceback
#
# # Setup logging
# logger = logging.getLogger(__name__)
#
# # Try to import PyMuPDF (fitz) for enhanced PDF processing
# try:
#     import fitz  # PyMuPDF
# except ImportError:
#     logger.warning("PyMuPDF (fitz) not installed. PDF link extraction disabled. Install with: pip install PyMuPDF")
#     fitz = None
#
#
# class ResumeParserService:
#     """
#     Service for parsing resumes and extracting structured data from them.
#     Uses user-provided API keys instead of global settings.
#     """
#
#     def __init__(self, user_openai_key=None, user_gemini_key=None):
#         """
#         Initialize the parser service with user-specific API keys.
#         Args:
#             user_openai_key (str, optional): User's OpenAI API key
#             user_gemini_key (str, optional): User's Google Gemini API key
#         """
#         self.user_openai_key = user_openai_key
#         self.user_gemini_key = user_gemini_key
#         self.openai_client = None
#         self.gemini_model = None
#
#         # Initialize clients if keys are provided
#         if self.user_openai_key:
#             try:
#                 self.openai_client = openai.OpenAI(api_key=self.user_openai_key)
#                 logger.info("OpenAI client initialized with user API key")
#             except Exception as e:
#                 logger.error(f"Failed to initialize OpenAI client: {str(e)}")
#
#         if self.user_gemini_key:
#             try:
#                 # Configure Gemini with user API key
#                 genai.configure(api_key=self.user_gemini_key)
#                 model_name = getattr(settings, 'GEMINI_MODEL_NAME', 'gemini-1.5-flash')
#                 self.gemini_model = genai.GenerativeModel(model_name)
#                 logger.info(f"Gemini model '{model_name}' initialized with user API key")
#             except Exception as e:
#                 logger.error(f"Failed to initialize Gemini model: {str(e)}")
#
#     def parse_resume(self, file_path, file_type=None, ai_parsing_enabled=True, ai_provider="gemini"):
#         """
#         Parse a resume file and extract structured data.
#
#         Args:
#             file_path (str): Path to the resume file
#             file_type (str, optional): File extension to override auto-detection
#             ai_parsing_enabled (bool): Whether to use AI for parsing
#             ai_provider (str): Which AI provider to use ("gemini" or "chatgpt")
#
#         Returns:
#             dict: Parsed resume data in structured format
#         """
#         try:
#             # Determine file type if not provided
#             if not file_type:
#                 file_type = os.path.splitext(file_path)[1].lower().strip('.')
#
#             # Extract text and links from the resume file
#             resume_text, extracted_links = self._extract_text_and_links(file_path, file_type)
#
#             # Check if text extraction was successful
#             if not resume_text or len(resume_text.strip()) < 50:
#                 logger.warning(f"Extracted text is too short ({len(resume_text)} chars). Extraction might have failed.")
#                 return {"error": "Could not extract sufficient text from the resume. Please check the file format."}
#
#             # Parse the resume text
#             if ai_parsing_enabled:
#                 if ai_provider == "chatgpt" and self.openai_client:
#                     parsed_data = self._parse_with_openai(resume_text, extracted_links)
#                 elif ai_provider == "gemini" and self.gemini_model:
#                     parsed_data = self._parse_with_gemini(resume_text, extracted_links)
#                 else:
#                     logger.warning(f"AI provider '{ai_provider}' not available or no API key. Using basic parsing.")
#                     parsed_data = self._basic_resume_parsing(resume_text)
#             else:
#                 logger.info("AI parsing disabled. Using basic parsing.")
#                 parsed_data = self._basic_resume_parsing(resume_text)
#
#             return parsed_data
#
#         except Exception as e:
#             logger.error(f"Error parsing resume: {str(e)}")
#             traceback.print_exc()
#             return {"error": f"Resume parsing failed: {str(e)}"}
#
#     def _extract_text_and_links(self, file_path, file_type):
#         """
#         Extract text and links from a resume file based on file type.
#
#         Args:
#             file_path (str): Path to the file
#             file_type (str): File extension (pdf, docx, doc, txt)
#
#         Returns:
#             tuple: (extracted_text, extracted_links)
#         """
#         if file_type == 'pdf':
#             return self._extract_from_pdf(file_path)
#         elif file_type in ['docx', 'doc']:
#             return self._extract_from_docx(file_path), []
#         elif file_type in ['txt', 'text', 'odt']:
#             return self._extract_from_txt(file_path), []
#         else:
#             logger.warning(f"Unsupported file type: {file_type}")
#             return "", []
#
#     def _extract_from_pdf(self, file_path):
#         """Extract text and links from a PDF file."""
#         text_content = ""
#         extracted_links = []
#
#         # Use PyMuPDF if available
#         if fitz:
#             doc = None
#             try:
#                 doc = fitz.open(file_path)
#                 for page_num in range(len(doc)):
#                     page = doc.load_page(page_num)
#                     text_content += page.get_text("text") + "\n"
#
#                     # Extract links from the page
#                     links = page.get_links()
#                     for link in links:
#                         if link.get('kind') == fitz.LINK_URI and 'uri' in link:
#                             uri = link['uri']
#                             if uri and uri not in extracted_links:
#                                 if isinstance(uri, str) and (uri.startswith('http') or uri.startswith('mailto:')):
#                                     extracted_links.append(uri)
#                                 elif isinstance(uri, str) and '@' in uri and '.' in uri and not uri.startswith(
#                                         'mailto:'):
#                                     if re.match(r"[^@]+@[^@]+\.[^@]+", uri):
#                                         extracted_links.append(f"mailto:{uri}")
#             except Exception as e:
#                 logger.error(f"Error extracting text/links from PDF with PyMuPDF: {str(e)}")
#                 return "", []
#             finally:
#                 if doc:
#                     try:
#                         doc.close()
#                     except Exception as e:
#                         logger.error(f"Error closing PDF document: {str(e)}")
#
#             logger.info(f"PyMuPDF extracted text length: {len(text_content)}, links: {len(extracted_links)}")
#             return text_content, extracted_links
#
#         # Fallback to PyPDF2 if PyMuPDF is not available
#         else:
#             try:
#                 import PyPDF2
#                 with open(file_path, 'rb') as file:
#                     pdf_reader = PyPDF2.PdfReader(file)
#
#                     # Check if PDF is encrypted
#                     if pdf_reader.is_encrypted:
#                         try:
#                             if pdf_reader.decrypt('') == 0:  # Failed decryption
#                                 logger.warning(f"Could not decrypt PDF {file_path}")
#                                 return "", []
#                         except Exception as e:
#                             logger.error(f"Error decrypting PDF {file_path}: {str(e)}")
#                             return "", []
#
#                     # Extract text from each page
#                     for page in pdf_reader.pages:
#                         page_text = page.extract_text()
#                         if page_text:
#                             text_content += page_text + "\n"
#
#                 logger.info(f"PyPDF2 extracted text length: {len(text_content)}")
#                 return text_content, []  # PyPDF2 doesn't extract links
#             except Exception as e:
#                 logger.error(f"Error during fallback PDF text extraction with PyPDF2: {str(e)}")
#                 return "", []
#
#     def _extract_from_docx(self, file_path):
#         """Extract text from a DOCX file."""
#         text = ""
#         try:
#             doc = docx.Document(file_path)
#             for para in doc.paragraphs:
#                 text += para.text + "\n"
#         except Exception as e:
#             logger.error(f"Error extracting text from DOCX: {str(e)}")
#         return text
#
#     def _extract_from_txt(self, file_path):
#         """Extract text from a plain text file."""
#         text = ""
#         try:
#             with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
#                 text = file.read()
#         except Exception as e:
#             logger.error(f"Error reading TXT file {file_path}: {str(e)}")
#             try:
#                 # Fallback to latin-1 encoding
#                 with open(file_path, 'r', encoding='latin-1', errors='ignore') as file:
#                     text = file.read()
#             except Exception as e:
#                 logger.error(f"Error reading TXT with latin-1: {str(e)}")
#         return text
#
#     def _parse_with_openai(self, resume_text, extracted_links):
#         """Parse resume text using OpenAI API."""
#         context_text = self._prepare_context_text(resume_text, extracted_links)
#         prompt = self._get_parsing_prompt(context_text)
#
#         try:
#             response = self.openai_client.chat.completions.create(
#                 model=getattr(settings, 'OPENAI_MODEL', 'gpt-4o-mini'),
#                 messages=[
#                     {"role": "system",
#                      "content": "Parse resume text into the specified JSON format. Ensure all values are correct types. List individual skills, do not categorize them. Output only valid JSON."},
#                     {"role": "user", "content": prompt}
#                 ],
#                 temperature=0.1,
#                 max_tokens=4000,
#                 response_format={"type": "json_object"}
#             )
#
#             result = response.choices[0].message.content
#
#             # Clean the response
#             result = re.sub(r'^```json\s*|\s*```$', '', result.strip(), flags=re.MULTILINE)
#             result = re.sub(r'^\s*//.*$', '', result, flags=re.MULTILINE)
#
#             # Parse the JSON
#             parsed_data = json.loads(result)
#
#             # Validate the structure
#             if not isinstance(parsed_data, dict):
#                 logger.error(f"Parsed data is not a dictionary (type: {type(parsed_data)})")
#                 return self._basic_resume_parsing(resume_text)
#
#             return parsed_data
#
#         except json.JSONDecodeError as e:
#             logger.error(f"JSON decode error from OpenAI: {str(e)}")
#             return self._basic_resume_parsing(resume_text)
#         except Exception as e:
#             logger.error(f"Error parsing with OpenAI: {str(e)}")
#             return self._basic_resume_parsing(resume_text)
#
#     def _parse_with_gemini(self, resume_text, extracted_links):
#         """Parse resume text using Google Gemini API."""
#         context_text = self._prepare_context_text(resume_text, extracted_links)
#         prompt = self._get_parsing_prompt(context_text)
#
#         try:
#             response = self.gemini_model.generate_content(
#                 prompt,
#                 generation_config=genai.types.GenerationConfig(
#                     temperature=0.2,
#                     response_mime_type="application/json"
#                 )
#             )
#
#             result = response.text
#
#             # Clean the response
#             result = re.sub(r'^```json\s*|\s*```$', '', result.strip(), flags=re.MULTILINE)
#             result = re.sub(r'^\s*//.*$', '', result, flags=re.MULTILINE)
#
#             # Parse the JSON
#             parsed_data = json.loads(result)
#
#             # Validate the structure
#             if not isinstance(parsed_data, dict):
#                 logger.error(f"Parsed data is not a dictionary (type: {type(parsed_data)})")
#                 return self._basic_resume_parsing(resume_text)
#
#             return parsed_data
#
#         except json.JSONDecodeError as e:
#             logger.error(f"JSON decode error from Gemini: {str(e)}")
#             return self._basic_resume_parsing(resume_text)
#         except Exception as e:
#             logger.error(f"Error parsing with Gemini: {str(e)}")
#             return self._basic_resume_parsing(resume_text)
#
#     def _prepare_context_text(self, resume_text, extracted_links):
#         """Prepare the context text with resume text and extracted links."""
#         context_text = resume_text
#
#         if extracted_links:
#             context_text += "\n\n--- Extracted Links (Prioritize these for URLs/Email) ---\n"
#             for link in extracted_links:
#                 context_text += f"- {link}\n"
#             context_text += "--- End Extracted Links ---\n"
#
#         return context_text
#
#     def _get_parsing_prompt(self, context_text):
#         """Get the prompt for resume parsing."""
#         return f"""
#         Extract information from the following resume text and format it STRICTLY as a valid JSON object.
#         Prioritize using URLs/Emails found in the 'Extracted Links' section if available and relevant, otherwise use URLs/Emails found directly in the text.
#         Use the EXACT keys specified below. Use `null` for optional fields that are missing, DO NOT omit the key. Ensure all string values are enclosed in double quotes.
#         **IMPORTANT for 'Skills': Extract each specific skill (like Python, SQL, Tableau, Leadership, Communication, AWS, Git) individually. DO NOT group skills into broad categories (like 'Programming Languages', 'Data Analytics', 'Software Development'). List every distinct skill found.**
#
#         JSON Structure:
#         {{
#           "Personal Information": {{
#             "First name": "string", "Middle name": "string | null", "Last name": "string",
#             "Email": "string (use mailto: links or email addresses found)", "Phone number": "string",
#             "Address": "string (e.g., 'City, State/Province') | null", "LinkedIn URL": "string (full URL) | null",
#             "GitHub URL": "string (full URL) | null", "Portfolio URL": "string (full URL) | null"
#           }},
#           "Professional Summary": "string | null",
#           "Skills": [ {{ "Skill name": "string (**Individual skill, e.g., 'Python', 'SQL', 'Communication', 'AWS' - NO CATEGORIES**)", "Skill type": "string (technical, soft, language, tool)", "Estimated proficiency level": "integer (0-100)" }} ] | null,
#           "Work Experience": [ {{ "Job title": "string", "Employer/Company name": "string", "Location": "string | null", "Start date": "string | null (YYYY-MM-DD)", "End date": "string | null (YYYY-MM-DD)", "Is current job": "boolean", "Bullet points": ["string", ...] | null }} ] | null,
#           "Education": [ {{ "School name": "string", "Location": "string | null", "Degree": "string", "Degree type": "string (high_school, associate, bachelor, master, doctorate, other)", "Field of study": "string | null", "Graduation date": "string | null (YYYY-MM-DD)", "GPA": "float | null" }} ] | null,
#           "Projects": [ {{ "Project name": "string", "Summary/description": "string | null", "Start date": "string | null (YYYY-MM-DD)", "Completion date": "string | null (YYYY-MM-DD)", "Project URL": "string | null", "GitHub URL": "string | null", "Bullet points": ["string", ...] | null }} ] | null,
#           "Certifications": [ {{ "Name": "string", "Institute/Issuing organization": "string | null", "Completion date": "string | null (YYYY-MM-DD)", "Expiration date": "string | null (YYYY-MM-DD)", "Score": "string | null", "URL/Link": "string | null", "Description": "string | null" }} ] | null,
#           "Languages": [ {{ "Language name": "string", "Proficiency": "string (basic, intermediate, advanced, native)" }} ] | null,
#           "Additional sections": [ {{ "Section name": "string", "Completion date": "string | null (YYYY-MM-DD)", "Bullet points": ["string", ...] | null, "Description": "string | null", "URL/Link": "string | null", "Institution name": "string | null" }} ] | null
#         }}
#
#         Ensure dates are YYYY-MM-DD format if possible, otherwise use text. Format Address/Location as "City, State/Province". Use key "Bullet points". Remember: List individual skills, not categories. Check Extracted Links section for relevant URLs/Emails. Output only valid JSON.
#
#         Resume Text & Links:
#         ---
#         {context_text}
#         ---
#         """
#
#     def _basic_resume_parsing(self, resume_text):
#         """Basic fallback parsing for when AI parsing fails."""
#         if not isinstance(resume_text, str):
#             resume_text = str(resume_text)
#
#         email = None
#         phone = None
#         linkedin = None
#         github = None
#         portfolio = None
#         first_name = ""
#         last_name = ""
#         middle_name = None
#
#         # Extract email
#         try:
#             email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
#             emails = re.findall(email_pattern, resume_text)
#             email = emails[0] if emails else None
#         except Exception as e:
#             logger.error(f"Error parsing email: {str(e)}")
#
#         # Extract phone
#         try:
#             phone_pattern = r'\b(?:\+\d{1,3}[-\.\s]?)?\(?\d{3}\)?[-\.\s]?\d{3}[-\.\s]?\d{4}\b'
#             phones = re.findall(phone_pattern, resume_text)
#             phone = phones[0] if phones else None
#         except Exception as e:
#             logger.error(f"Error parsing phone: {str(e)}")
#
#         # Extract URLs
#         try:
#             url_pattern = r'https?://(?:www\.)?[-a-zA-Z0-9@:%._\+~#=]{1,256}\.[a-zA-Z0-9()]{1,6}\b(?:[-a-zA-Z0-9()@:%_\+.~#?&//=]*)'
#             urls = re.findall(url_pattern, resume_text)
#
#             for url in urls:
#                 try:
#                     parsed = urlparse(url)
#                     scheme = parsed.scheme or 'https'
#                     netloc = parsed.netloc.replace('www.', '')
#                     path = parsed.path or ''
#                     clean_url = f"{scheme}://{netloc}{path}".rstrip('/')
#
#                     if 'linkedin.com/in/' in clean_url and not linkedin:
#                         linkedin = clean_url
#                     elif 'github.com/' in clean_url and not github:
#                         github = clean_url
#                     elif not any(x in clean_url for x in ['linkedin.com', 'github.com']) and not portfolio:
#                         if not any(ext in clean_url for ext in ['.pdf', '.png', '.jpg', '.jpeg', '.gif']):
#                             portfolio = clean_url
#                 except ValueError:
#                     logger.warning(f"Skipping malformed URL: {url}")
#                     continue
#         except Exception as e:
#             logger.error(f"Error parsing URLs: {str(e)}")
#
#         # Extract name
#         try:
#             lines = resume_text.strip().split('\n')
#             name = lines[0].strip() if lines else ""
#             name_parts = name.split()
#             first_name = name_parts[0] if name_parts else ""
#             last_name = name_parts[-1] if len(name_parts) > 1 else ""
#             middle_name = " ".join(name_parts[1:-1]) if len(name_parts) > 2 else None
#         except Exception as e:
#             logger.error(f"Error parsing name: {str(e)}")
#
#         # Return the default dictionary structure
#         return {
#             "Personal Information": {
#                 "First name": first_name or None,
#                 "Middle name": middle_name,
#                 "Last name": last_name or None,
#                 "Email": email,
#                 "Phone number": phone,
#                 "Address": None,
#                 "LinkedIn URL": linkedin,
#                 "GitHub URL": github,
#                 "Portfolio URL": portfolio
#             },
#             "Professional Summary": None,
#             "Skills": [],
#             "Work Experience": [],
#             "Education": [],
#             "Projects": [],
#             "Certifications": [],
#             "Languages": [],
#             "Additional sections": []
#         }
#
#     @staticmethod
#     def format_date(date_str):
#         """Convert various date string formats to YYYY-MM-DD."""
#         if not date_str or not isinstance(date_str, str) or str(date_str).lower() == 'present':
#             return None
#
#         date_str = str(date_str).strip()
#
#         # Try common date formats
#         formats = [
#             '%Y-%m-%d', '%m/%d/%Y', '%m-%d-%Y', '%d-%b-%Y', '%d-%B-%Y',
#             '%Y/%m/%d', '%d %b %Y', '%d %B %Y', '%b %d, %Y', '%B %d, %Y',
#             '%Y-%m', '%m/%Y', '%b %Y', '%B %Y', '%Y',
#         ]
#
#         for fmt in formats:
#             try:
#                 dt = datetime.strptime(date_str, fmt)
#
#                 # Handle formats with varying specificity
#                 if fmt in ['%Y-%m', '%m/%Y', '%b %Y', '%B %Y']:
#                     return dt.strftime('%Y-%m-01')  # Default to the 1st of the month
#                 elif fmt == '%Y':
#                     return dt.strftime('%Y-01-01')  # Default to Jan 1st of the year
#                 else:
#                     return dt.strftime('%Y-%m-%d')  # Full date
#             except ValueError:
#                 continue
#
#         # Fallback attempts
#         try:
#             # Try to match "Month Year" format
#             match_month_year = re.match(r'([A-Za-z]{3,})\s+(\d{4})', date_str)
#             if match_month_year:
#                 month_str, year_str = match_month_year.groups()
#
#                 try:
#                     dt = datetime.strptime(f"{month_str} 1 {year_str}", '%b 1 %Y')
#                     return dt.strftime('%Y-%m-01')
#                 except ValueError:
#                     try:
#                         dt = datetime.strptime(f"{month_str} 1 {year_str}", '%B 1 %Y')
#                         return dt.strftime('%Y-%m-01')
#                     except ValueError:
#                         pass
#
#             # Try to match just a year
#             match_year = re.match(r'(\d{4})', date_str)
#             if match_year:
#                 year = match_year.group(1)
#                 return f"{year}-01-01"
#         except Exception as e:
#             logger.error(f"Error during fallback date parsing for '{date_str}': {str(e)}")
#
#         logger.warning(f"Could not parse date string: {date_str}")
#         return None
#
#     @staticmethod
#     def format_url(url_str):
#         """Adds https:// if missing and cleans the URL."""
#         if not url_str or not isinstance(url_str, str):
#             return None
#
#         url_str = url_str.strip()
#         if not url_str:
#             return None
#
#         try:
#             # Add scheme if missing, default to https
#             if not re.match(r"^(?:[a-z]+:)?//", url_str, re.IGNORECASE) and '.' in url_str:
#                 url_str = 'https://' + url_str
#
#             parsed = urlparse(url_str)
#             scheme = parsed.scheme
#             netloc = parsed.netloc
#             path = parsed.path or ''
#
#             # Basic validation
#             if not scheme or not netloc or '.' not in netloc:
#                 # Check if it's an email
#                 if '@' in url_str and '.' in url_str and not url_str.startswith('mailto:'):
#                     if re.match(r"[^@]+@[^@]+\.[^@]+", url_str):
#                         return f"mailto:{url_str}"
#                 return None
#
#             # Reconstruct URL, removing trailing slash if path is not just '/'
#             clean_url = f"{scheme}://{netloc}{path}"
#             if clean_url.endswith('/') and len(path) > 1:
#                 clean_url = clean_url.rstrip('/')
#
#             return clean_url
#         except Exception as e:
#             logger.error(f"Error formatting URL '{url_str}': {str(e)}")
#             return None
#
#     @staticmethod
#     def format_location(location_data):
#         """Formats location data into a string."""
#         if isinstance(location_data, dict):
#             city = location_data.get('City', '')
#             state = location_data.get('State/Province', '')
#             country = location_data.get('Country', '')
#
#             parts = [str(part).strip() for part in [city, state, country]
#                      if part and isinstance(part, str) and str(part).strip()]
#             return ", ".join(parts) if parts else ""
#
#         elif isinstance(location_data, str):
#             return location_data.strip()
#
#         elif isinstance(location_data, (list, tuple)):
#             parts = [str(item).strip() for item in location_data
#                      if item and isinstance(item, str) and str(item).strip()]
#             return ", ".join(parts) if parts else ""
#
#         else:
#             return ""
#
#     @staticmethod
#     def safe_strip(value, default=''):
#         """Safely strips a value if it's a string, joins if list/tuple, returns default otherwise."""
#         if isinstance(value, str):
#             return value.strip()
#
#         elif isinstance(value, (list, tuple)):
#             string_items = [str(item).strip() for item in value
#                             if isinstance(item, str) and str(item).strip()]
#             return " ".join(string_items) if string_items else default
#
#         elif value is None:
#             return default
#
#         else:
#             try:
#                 return str(value).strip()
#             except:
#                 return default
#
# # # File: services/parser/resume_parser_service.py
# #
# #
# # # services/resume_parser_service.py
# #
# # import os
# # import uuid # Import uuid
# # import tempfile
# # # import PyPDF2 # Only imported as fallback now
# # import docx
# # import re
# # import json
# # from django.conf import settings
# # import openai
# # import google.generativeai as genai
# # from django.utils import timezone
# # from datetime import datetime
# # from urllib.parse import urlparse # Import urlparse
# # import traceback # For detailed error logging
# #
# # # Import PyMuPDF (fitz) - INSTALL IT FIRST: pip install PyMuPDF
# # try:
# #     import fitz # PyMuPDF
# # except ImportError:
# #     print("PyMuPDF (fitz) not installed. PDF link extraction disabled. Install with: pip install PyMuPDF")
# #     fitz = None
# #
# #
# # # --- resume_upload_path ---
# # def resume_upload_path(instance, filename):
# #     """Generate a unique path for uploaded resume files."""
# #     ext = filename.split('.')[-1]
# #     filename = f"{uuid.uuid4()}.{ext}"
# #     # Ensure instance.user and instance.user.id exist before using them
# #     user_id = getattr(getattr(instance, 'user', None), 'id', 'unknown_user')
# #     return os.path.join(
# #         'resumes', f"user_{user_id}", str(timezone.now().year),
# #         str(timezone.now().month), filename
# #     )
# #
# # # Initialize APIs
# # openai.api_key = getattr(settings, 'OPENAI_API_KEY', None)
# # genai_api_key = getattr(settings, 'GOOGLE_GENAI_API_KEY', None)
# # if genai_api_key:
# #     try:
# #         genai.configure(api_key=genai_api_key)
# #         print("Google Generative AI configured successfully.")
# #     except Exception as e:
# #         print(f"Error configuring Google Generative AI: {e}")
# #         genai_api_key = None # Mark as not configured on error
# # else:
# #     print("Warning: Google Generative AI API key not configured in settings.")
# #
# # if not openai.api_key:
# #     print("Warning: OpenAI API key not configured in settings.")
# #
# #
# # # --- Text Extraction Functions ---
# # def extract_text_and_links_from_pdf(file_path):
# #     """Extract text content AND hyperlink URIs from a PDF file using PyMuPDF."""
# #     text_content = ""
# #     extracted_links = []
# #     if not fitz: # Check if PyMuPDF is installed
# #         print("PyMuPDF not found, falling back to basic text extraction for PDF.")
# #         try:
# #             # Fallback using PyPDF2 (consider adding PyPDF2 to requirements if needed)
# #             import PyPDF2
# #             with open(file_path, 'rb') as file:
# #                 pdf_reader = PyPDF2.PdfReader(file)
# #                 if pdf_reader.is_encrypted:
# #                     try:
# #                         # Attempt decryption with an empty password
# #                         if pdf_reader.decrypt('') == 0: # 0 indicates failure
# #                              print(f"Warning: Could not decrypt PDF {file_path}")
# #                              return "", [] # Return empty if decryption fails
# #                     except Exception as decrypt_err:
# #                          print(f"Error decrypting PDF {file_path}: {decrypt_err}")
# #                          return "", [] # Return empty on decryption error
# #
# #                 text_content = ""
# #                 for page in pdf_reader.pages:
# #                     page_text = page.extract_text()
# #                     if page_text: # Check if text extraction was successful
# #                         text_content += page_text + "\n" # Add newline between pages
# #                 print(f"PyPDF2 extracted text length: {len(text_content)}")
# #             return text_content, []
# #         except Exception as e_pypdf:
# #             print(f"Error during fallback PDF text extraction with PyPDF2: {e_pypdf}")
# #             return "", []
# #
# #     doc = None # Initialize doc to None
# #     try:
# #         doc = fitz.open(file_path)
# #         for page_num in range(len(doc)):
# #             page = doc.load_page(page_num)
# #             text_content += page.get_text("text") + "\n" # Get plain text
# #             links = page.get_links()
# #             for link in links:
# #                 # Check if 'uri' key exists and kind is URI link
# #                 if link.get('kind') == fitz.LINK_URI and 'uri' in link:
# #                     uri = link['uri']
# #                     if uri and uri not in extracted_links:
# #                         # Basic validation for URI format
# #                         if isinstance(uri, str) and (uri.startswith('http') or uri.startswith('mailto:')):
# #                             extracted_links.append(uri)
# #                         elif isinstance(uri, str) and '@' in uri and '.' in uri and not uri.startswith('mailto:'):
# #                             # Attempt to fix potential email links missing mailto:
# #                              if re.match(r"[^@]+@[^@]+\.[^@]+", uri): # Basic email format check
# #                                 extracted_links.append(f"mailto:{uri}")
# #     except Exception as e_fitz:
# #         print(f"Error extracting text/links from PDF with PyMuPDF: {e_fitz}")
# #         # If PyMuPDF fails, potentially fallback to PyPDF2 again, or return empty
# #         return "", [] # Return empty on PyMuPDF error
# #     finally:
# #         if doc:
# #             try: doc.close()
# #             except Exception as e_close: print(f"Error closing PDF document: {e_close}")
# #
# #     print(f"PyMuPDF extracted text length: {len(text_content)}")
# #     print(f"Extracted Links: {extracted_links}")
# #     return text_content, extracted_links
# #
# # def extract_text_from_docx(file_path):
# #     text = "";
# #     try:
# #         doc = docx.Document(file_path)
# #         for para in doc.paragraphs: text += para.text + "\n"
# #     except Exception as e: print(f"Error extracting text from DOCX: {str(e)}")
# #     return text
# #
# # def extract_text_from_txt(file_path):
# #     text = "";
# #     try:
# #         # Try UTF-8 first, ignore errors
# #         with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
# #              text = file.read()
# #     except Exception as e:
# #         print(f"Error reading TXT file {file_path}: {str(e)}")
# #         # Optionally try other encodings here if needed, e.g., latin-1
# #         # try:
# #         #     with open(file_path, 'r', encoding='latin-1', errors='ignore') as file:
# #         #         text = file.read()
# #         # except Exception as e_latin:
# #         #     print(f"Error reading TXT with latin-1: {e_latin}")
# #     return text
# #
# # def extract_text_from_resume(file_obj):
# #     """Extract text and potentially links from various resume file formats."""
# #     if not file_obj or not hasattr(file_obj, 'name'):
# #         print("Error: Invalid file object passed to extract_text_from_resume.")
# #         # Return empty string and empty list for consistency
# #         return "", []
# #
# #     # Ensure file object is reset if it has been read before
# #     if hasattr(file_obj, 'seek'):
# #         try:
# #             file_obj.seek(0)
# #         except Exception as e:
# #             print(f"Warning: Could not seek file object - {e}")
# #
# #     file_path = None # Initialize file_path
# #     suffix = os.path.splitext(file_obj.name)[1].lower()
# #     # Use a 'with' statement for better file handling
# #     try:
# #         # Create temp file explicitly to manage deletion better
# #         fd, file_path = tempfile.mkstemp(suffix=suffix)
# #         os.close(fd) # Close the file descriptor
# #
# #         with open(file_path, 'wb') as temp_file_write:
# #              for chunk in file_obj.chunks():
# #                  temp_file_write.write(chunk)
# #
# #         text = ""; links = []
# #         file_ext = suffix
# #         if file_ext == '.pdf':
# #             text, links = extract_text_and_links_from_pdf(file_path) # Use new function
# #         elif file_ext in ['.docx', '.doc']:
# #             text = extract_text_from_docx(file_path)
# #         elif file_ext in ['.txt', '.odt']: # Consider adding specific ODT handling if needed
# #             text = extract_text_from_txt(file_path)
# #         else:
# #             print(f"Unsupported file format: {file_ext}")
# #         # Ensure text is always a string, even if extraction fails
# #         if not isinstance(text, str):
# #              text = ""
# #
# #     except Exception as e:
# #         print(f"Error during text extraction setup or file write: {e}")
# #         text = ""
# #         links = []
# #     finally:
# #         # Ensure temporary file is deleted
# #         if file_path and os.path.exists(file_path):
# #             try:
# #                 os.unlink(file_path)
# #             except Exception as e_unlink:
# #                  print(f"Error deleting temporary file {file_path}: {e_unlink}")
# #
# #     # Add a check for very short text which might indicate an extraction issue
# #     if len(text) < 50: # Adjusted threshold
# #          print(f"Warning: Extracted text is very short ({len(text)} chars). Extraction might have failed.")
# #
# #     return text, links # Return both text and links
# #
# #
# # # --- AI Parsing Function (Prompt Adjusted for Links) ---
# # def parse_resume_with_ai(resume_text, extracted_links=None, ai_engine='chatgpt'):
# #     """Parse resume text using AI to extract structured data, considering extracted links."""
# #     # Add a check for minimal text length before attempting AI parsing
# #     if not resume_text or len(resume_text.strip()) < 50: # Use same threshold as extraction
# #         print("Error: Resume text is too short or empty for AI parsing. Falling back to basic parsing.")
# #         return basic_resume_parsing(resume_text) # Return default dict
# #
# #     # Default to empty list if None
# #     if extracted_links is None:
# #         extracted_links = []
# #
# #     context_text = resume_text
# #     if extracted_links:
# #         context_text += "\n\n--- Extracted Links (Prioritize these for URLs/Email) ---\n"
# #         for link in extracted_links: context_text += f"- {link}\n"
# #         context_text += "--- End Extracted Links ---\n"
# #
# #     # --- REFINED PROMPT ---
# #     # *** Includes explicit instructions for Skills ***
# #     prompt = f"""
# #     Extract information from the following resume text and format it STRICTLY as a valid JSON object.
# #     Prioritize using URLs/Emails found in the 'Extracted Links' section if available and relevant, otherwise use URLs/Emails found directly in the text.
# #     Use the EXACT keys specified below. Use `null` for optional fields that are missing, DO NOT omit the key. Ensure all string values are enclosed in double quotes.
# #     **IMPORTANT for 'Skills': Extract each specific skill (like Python, SQL, Tableau, Leadership, Communication, AWS, Git) individually. DO NOT group skills into broad categories (like 'Programming Languages', 'Data Analytics', 'Software Development'). List every distinct skill found.**
# #
# #     JSON Structure:
# #     {{
# #       "Personal Information": {{
# #         "First name": "string", "Middle name": "string | null", "Last name": "string",
# #         "Email": "string (use mailto: links or email addresses found)", "Phone number": "string",
# #         "Address": "string (e.g., 'City, State/Province') | null", "LinkedIn URL": "string (full URL) | null",
# #         "GitHub URL": "string (full URL) | null", "Portfolio URL": "string (full URL) | null"
# #       }},
# #       "Professional Summary": "string | null",
# #       "Skills": [ {{ "Skill name": "string (**Individual skill, e.g., 'Python', 'SQL', 'Communication', 'AWS' - NO CATEGORIES**)", "Skill type": "string (technical, soft, language, tool)", "Estimated proficiency level": "integer (0-100)" }} ] | null,
# #       "Work Experience": [ {{ "Job title": "string", "Employer/Company name": "string", "Location": "string | null", "Start date": "string | null (YYYY-MM-DD)", "End date": "string | null (YYYY-MM-DD)", "Is current job": "boolean", "Bullet points": ["string", ...] | null }} ] | null,
# #       "Education": [ {{ "School name": "string", "Location": "string | null", "Degree": "string", "Degree type": "string (high_school, associate, bachelor, master, doctorate, other)", "Field of study": "string | null", "Graduation date": "string | null (YYYY-MM-DD)", "GPA": "float | null" }} ] | null,
# #       "Projects": [ {{ "Project name": "string", "Summary/description": "string | null", "Start date": "string | null (YYYY-MM-DD)", "Completion date": "string | null (YYYY-MM-DD)", "Project URL": "string | null", "GitHub URL": "string | null", "Bullet points": ["string", ...] | null }} ] | null,
# #       "Certifications": [ {{ "Name": "string", "Institute/Issuing organization": "string | null", "Completion date": "string | null (YYYY-MM-DD)", "Expiration date": "string | null (YYYY-MM-DD)", "Score": "string | null", "URL/Link": "string | null", "Description": "string | null" }} ] | null,
# #       "Languages": [ {{ "Language name": "string", "Proficiency": "string (basic, intermediate, advanced, native)" }} ] | null,
# #       "Additional sections": [ {{ "Section name": "string", "Completion date": "string | null (YYYY-MM-DD)", "Bullet points": ["string", ...] | null, "Description": "string | null", "URL/Link": "string | null", "Institution name": "string | null" }} ] | null
# #     }}
# #
# #     Ensure dates are YYYY-MM-DD format if possible, otherwise use text. Format Address/Location as "City, State/Province". Use key "Bullet points". Remember: List individual skills, not categories. Check Extracted Links section for relevant URLs/Emails. Output only valid JSON.
# #
# #     Resume Text & Links:
# #     ---
# #     {context_text}
# #     ---
# #     """
# #
# #     result = None
# #     try:
# #         # Clear previous errors if any (optional, depends on how errors are handled)
# #         # clear_previous_parsing_error()
# #
# #         if ai_engine == 'chatgpt' and openai.api_key:
# #             print("--- Using ChatGPT for parsing ---")
# #             response = openai.chat.completions.create(
# #                 model=getattr(settings, 'OPENAI_MODEL', 'gpt-4o-mini'), # Use a capable model
# #                 messages=[
# #                     {"role": "system", "content": "Parse resume text into the specified JSON format. Ensure all values are correct types (string, number, boolean, null, array). List individual skills, do not categorize them. Prioritize URLs/Emails from 'Extracted Links' section. Output only valid JSON."}, # Added instruction here too
# #                     {"role": "user", "content": prompt}
# #                 ],
# #                 temperature=0.1,
# #                 max_tokens=4000, # Increase max tokens if resumes are long
# #                 response_format={ "type": "json_object" } # Request JSON output format
# #             )
# #             result = response.choices[0].message.content
# #         elif ai_engine == 'gemini' and genai_api_key:
# #             print("--- Using Gemini for parsing ---")
# #             ai_settings = getattr(settings, 'AI_SETTINGS', {})
# #             gemini_config = ai_settings.get('gemini', {})
# #             model_name = gemini_config.get('model', 'gemini-1.5-flash') # Use Flash or Pro
# #             model = genai.GenerativeModel(model_name)
# #             # Add system instruction for Gemini if model supports it (check Gemini docs)
# #             # Example (conceptual):
# #             # system_instruction = "Parse resume text into the specified JSON format... List individual skills, do not categorize them."
# #             # response = model.generate_content([system_instruction, prompt], ...)
# #             response = model.generate_content(
# #                 prompt, # Sending the main prompt which now has the instruction
# #                 generation_config=genai.types.GenerationConfig(
# #                     temperature=0.2,
# #                     response_mime_type="application/json" # Request JSON output
# #                 )
# #                 # safety_settings=... # Consider adding safety settings
# #             )
# #             result = response.text
# #         else:
# #             print("Warning: No valid AI engine configured/API key missing. Falling back to basic parsing.")
# #             return basic_resume_parsing(resume_text)
# #
# #         if not result:
# #             print("Error: AI returned an empty response. Falling back.")
# #             return basic_resume_parsing(resume_text)
# #
# #         # Clean the response (remove potential markdown/comments)
# #         result = re.sub(r'^```json\s*|\s*```$', '', result.strip(), flags=re.MULTILINE)
# #         result = re.sub(r'^\s*//.*$', '', result, flags=re.MULTILINE) # Remove potential comments
# #
# #         # Attempt to load the JSON
# #         parsed_data = json.loads(result)
# #         print("--- AI Parsing Successful ---")
# #         # print(json.dumps(parsed_data, indent=2)) # Optional: Pretty print parsed data
# #         # **Crucial Check:** Ensure the top-level structure is a dictionary
# #         if not isinstance(parsed_data, dict):
# #             print(f"Error: Parsed data is not a dictionary after JSON load (type: {type(parsed_data)}). Falling back.")
# #             return basic_resume_parsing(resume_text)
# #
# #         return parsed_data
# #
# #     except json.JSONDecodeError as json_err:
# #         print(f"Error decoding JSON response from AI: {json_err}")
# #         print(f"AI Response causing error:\n>>>\n{result}\n<<<")
# #         # Maybe try to fix common JSON errors here or fallback
# #         return basic_resume_parsing(resume_text) # Fallback on JSON error
# #     except Exception as e:
# #         print(f"Error parsing resume with AI ({ai_engine}): {str(e)}")
# #         traceback.print_exc()
# #         return basic_resume_parsing(resume_text) # Fallback on other errors
# #
# #
# # # --- Basic Fallback Parser (Returns Dict) ---
# # def basic_resume_parsing(resume_text):
# #     """Basic fallback parsing function if AI parsing fails. Returns dictionary."""
# #     # Ensure resume_text is a string
# #     if not isinstance(resume_text, str):
# #         resume_text = str(resume_text) # Attempt to convert
# #
# #     email=None; phone=None; linkedin=None; github=None; portfolio=None; first_name=""; last_name=""; middle_name=None
# #     try:
# #         email_pattern=r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'; emails=re.findall(email_pattern, resume_text); email=emails[0] if emails else None
# #     except Exception as e: print(f"Error parsing email: {e}")
# #     try:
# #         phone_pattern=r'\b(?:\+\d{1,3}[-\.\s]?)?\(?\d{3}\)?[-\.\s]?\d{3}[-\.\s]?\d{4}\b'; phones=re.findall(phone_pattern, resume_text); phone=phones[0] if phones else None
# #     except Exception as e: print(f"Error parsing phone: {e}")
# #     try:
# #         url_pattern=r'https?://(?:www\.)?[-a-zA-Z0-9@:%._\+~#=]{1,256}\.[a-zA-Z0-9()]{1,6}\b(?:[-a-zA-Z0-9()@:%_\+.~#?&//=]*)'; urls=re.findall(url_pattern, resume_text)
# #         for url in urls:
# #             try:
# #                 parsed=urlparse(url); scheme=parsed.scheme or 'https'; netloc=parsed.netloc.replace('www.',''); path=parsed.path or ''; clean_url=f"{scheme}://{netloc}{path}".rstrip('/') # Ensure path is string
# #                 if 'linkedin.com/in/' in clean_url and not linkedin:
# #                     linkedin=clean_url
# #                 elif 'github.com/' in clean_url and not github:
# #                     github=clean_url
# #                 elif not any(x in clean_url for x in ['linkedin.com','github.com']) and not portfolio:
# #                     # Basic check to avoid common non-portfolio URLs
# #                     if not any(ext in clean_url for ext in ['.pdf', '.png', '.jpg', '.jpeg', '.gif']):
# #                         portfolio=clean_url
# #             except ValueError:
# #                  print(f"Skipping malformed URL in basic_resume_parsing: {url}")
# #                  continue
# #     except Exception as e: print(f"Error parsing URLs: {e}")
# #     try:
# #         lines=resume_text.strip().split('\n'); name=lines[0].strip() if lines else ""
# #         name_parts=name.split(); first_name=name_parts[0] if name_parts else ""; last_name=name_parts[-1] if len(name_parts)>1 else ""; middle_name=" ".join(name_parts[1:-1]) if len(name_parts)>2 else None
# #     except Exception as e: print(f"Error parsing name: {e}")
# #
# #     # Return the default dictionary structure
# #     # Ensure all top-level keys are present even if empty
# #     return {
# #         "Personal Information": {
# #             "First name": first_name or None,
# #             "Middle name": middle_name,
# #             "Last name": last_name or None,
# #             "Email": email,
# #             "Phone number": phone,
# #             "Address": None,
# #             "LinkedIn URL": linkedin,
# #             "GitHub URL": github,
# #             "Portfolio URL": portfolio
# #         },
# #         "Professional Summary": None,
# #         "Skills": [], # Use empty list instead of None for consistency
# #         "Work Experience": [],
# #         "Education": [],
# #         "Projects": [],
# #         "Certifications": [],
# #         "Languages": [],
# #         "Additional sections": []
# #     }
# #
# #
# # # --- Helper Functions ---
# # def format_date(date_str):
# #     """Convert various date string formats to YYYY-MM-DD."""
# #     if not date_str or not isinstance(date_str, str) or str(date_str).lower() == 'present':
# #         return None
# #     date_str = str(date_str).strip()
# #     # Order formats from more specific to less specific
# #     formats = [
# #         '%Y-%m-%d', '%m/%d/%Y', '%m-%d-%Y', '%d-%b-%Y', '%d-%B-%Y',
# #         '%Y/%m/%d', '%d %b %Y', '%d %B %Y', '%b %d, %Y', '%B %d, %Y',
# #         '%Y-%m', '%m/%Y', '%b %Y', '%B %Y', '%Y',
# #     ]
# #     for fmt in formats:
# #         try:
# #             dt = datetime.strptime(date_str, fmt)
# #             # Handle formats that only provide year and month or just year
# #             if fmt in ['%Y-%m', '%m/%Y', '%b %Y', '%B %Y']:
# #                 return dt.strftime('%Y-%m-01') # Default to the 1st of the month
# #             elif fmt == '%Y':
# #                 return dt.strftime('%Y-01-01') # Default to Jan 1st of the year
# #             else:
# #                 return dt.strftime('%Y-%m-%d') # Full date
# #         except ValueError:
# #             continue # Try the next format
# #
# #     # Fallback attempts (keep simple for now)
# #     try:
# #         match_month_year = re.match(r'([A-Za-z]{3,})\s+(\d{4})', date_str)
# #         if match_month_year:
# #             month_str, year_str = match_month_year.groups()
# #             # Try common month abbreviations/names
# #             try:
# #                 dt = datetime.strptime(f"{month_str} 1 {year_str}", '%b 1 %Y')
# #                 return dt.strftime('%Y-%m-01')
# #             except ValueError:
# #                 try:
# #                     dt = datetime.strptime(f"{month_str} 1 {year_str}", '%B 1 %Y')
# #                     return dt.strftime('%Y-%m-01')
# #                 except ValueError: pass # Give up on month/year text
# #
# #         # Attempt to parse just YYYY as final fallback
# #         match_year = re.match(r'(\d{4})', date_str)
# #         if match_year:
# #             year = match_year.group(1)
# #             return datetime(int(year), 1, 1).strftime('%Y-01-01')
# #
# #     except Exception as e:
# #          print(f"Error during fallback date parsing for '{date_str}': {e}")
# #
# #     print(f"Warning: Could not parse date string: {date_str}")
# #     return None # Return None if no format matches
# #
# # def format_url(url_str):
# #     """Adds https:// if missing and cleans the URL."""
# #     if not url_str or not isinstance(url_str, str): return None; url_str = url_str.strip();
# #     if not url_str: return None;
# #     try:
# #         # Add scheme if missing, default to https
# #         if not re.match(r"^(?:[a-z]+:)?//", url_str, re.IGNORECASE) and '.' in url_str:
# #              url_str = 'https://' + url_str
# #
# #         parsed = urlparse(url_str);
# #         scheme = parsed.scheme; netloc = parsed.netloc; path = parsed.path or ''; # Ensure path is string
# #
# #         # Basic validation
# #         if not scheme or not netloc or '.' not in netloc:
# #             # If it looks like an email, add mailto:
# #             if '@' in url_str and '.' in url_str and not url_str.startswith('mailto:'):
# #                 if re.match(r"[^@]+@[^@]+\.[^@]+", url_str): # Basic email format check
# #                     return f"mailto:{url_str}"
# #             return None # Invalid URL structure
# #
# #         # Reconstruct URL, remove trailing slash if path is more than just '/'
# #         clean_url = f"{scheme}://{netloc}{path}"
# #         if clean_url.endswith('/') and len(path) > 1:
# #             clean_url = clean_url.rstrip('/')
# #         return clean_url
# #     except Exception as e: print(f"Error formatting URL '{url_str}': {e}"); return None
# #
# # def format_location(location_data):
# #     """Formats location data (string, dict, list, tuple) into a string."""
# #     if isinstance(location_data, dict):
# #         city = location_data.get('City', '')
# #         state = location_data.get('State/Province', '')
# #         country = location_data.get('Country', '')
# #         # Filter out None or empty strings before joining
# #         parts = [str(part).strip() for part in [city, state, country] if part and isinstance(part, str) and str(part).strip()]
# #         return ", ".join(parts) if parts else ""
# #     elif isinstance(location_data, str):
# #         return location_data.strip()
# #     elif isinstance(location_data, (list, tuple)): # Handle list/tuple case
# #         # Join elements if they are strings, filter out non-strings/empty
# #         parts = [str(item).strip() for item in location_data if item and isinstance(item, str) and str(item).strip()]
# #         return ", ".join(parts) if parts else ""
# #     else:
# #         # Return empty string for other types (like None, int, float, etc.)
# #         return ""
# #
# # def safe_strip(value, default=''):
# #     """Safely strips a value if it's a string, joins if list/tuple, otherwise returns default."""
# #     if isinstance(value, str):
# #         return value.strip()
# #     elif isinstance(value, (list, tuple)):
# #         # Join elements with a space if it's a list/tuple of strings
# #         # Filter out non-string items before joining
# #         string_items = [str(item).strip() for item in value if isinstance(item, str) and str(item).strip()]
# #         return " ".join(string_items) if string_items else default
# #     # Handle cases where value might be None, int, float, etc.
# #     elif value is None:
# #         return default
# #     else:
# #         # Attempt to convert other types to string, then return default if fails
# #         try:
# #             return str(value).strip()
# #         except:
# #             return default
# #
# # # # --- create_resume_from_parsed_data (Obsolete - Remove if unused or implement) ---
# # # # def create_resume_from_parsed_data(user, parsed_data, template_id=None):
# # # #     # This function seems to be defined elsewhere or needs implementation
# # # #     pass
# # #
# # # # # services/resume_parser_service.py
# # # #
# # # # import os
# # # # import uuid # Import uuid
# # # # import tempfile
# # # # # import PyPDF2 # Only imported as fallback now
# # # # import docx
# # # # import re
# # # # import json
# # # # from django.conf import settings
# # # # import openai
# # # # import google.generativeai as genai
# # # # from django.utils import timezone
# # # # from datetime import datetime
# # # # from urllib.parse import urlparse # Import urlparse
# # # # import traceback # For detailed error logging
# # # #
# # # # # Import PyMuPDF (fitz) - INSTALL IT FIRST: pip install PyMuPDF
# # # # try:
# # # #     import fitz # PyMuPDF
# # # # except ImportError:
# # # #     print("PyMuPDF (fitz) not installed. PDF link extraction disabled. Install with: pip install PyMuPDF")
# # # #     fitz = None
# # # #
# # # #
# # # # # --- resume_upload_path ---
# # # # def resume_upload_path(instance, filename):
# # # #     """Generate a unique path for uploaded resume files."""
# # # #     ext = filename.split('.')[-1]
# # # #     filename = f"{uuid.uuid4()}.{ext}"
# # # #     # Ensure instance.user and instance.user.id exist before using them
# # # #     user_id = getattr(getattr(instance, 'user', None), 'id', 'unknown_user')
# # # #     return os.path.join(
# # # #         'resumes', f"user_{user_id}", str(timezone.now().year),
# # # #         str(timezone.now().month), filename
# # # #     )
# # # #
# # # # # Initialize APIs
# # # # openai.api_key = getattr(settings, 'OPENAI_API_KEY', None)
# # # # genai_api_key = getattr(settings, 'GOOGLE_GENAI_API_KEY', None)
# # # # if genai_api_key:
# # # #     try:
# # # #         genai.configure(api_key=genai_api_key)
# # # #         print("Google Generative AI configured successfully.")
# # # #     except Exception as e:
# # # #         print(f"Error configuring Google Generative AI: {e}")
# # # #         genai_api_key = None # Mark as not configured on error
# # # # else:
# # # #     print("Warning: Google Generative AI API key not configured in settings.")
# # # #
# # # # if not openai.api_key:
# # # #     print("Warning: OpenAI API key not configured in settings.")
# # # #
# # # #
# # # # # --- Text Extraction Functions ---
# # # # def extract_text_and_links_from_pdf(file_path):
# # # #     """Extract text content AND hyperlink URIs from a PDF file using PyMuPDF."""
# # # #     text_content = ""
# # # #     extracted_links = []
# # # #     if not fitz: # Check if PyMuPDF is installed
# # # #         print("PyMuPDF not found, falling back to basic text extraction for PDF.")
# # # #         try:
# # # #             # Fallback using PyPDF2 (consider adding PyPDF2 to requirements if needed)
# # # #             import PyPDF2
# # # #             with open(file_path, 'rb') as file:
# # # #                 pdf_reader = PyPDF2.PdfReader(file)
# # # #                 if pdf_reader.is_encrypted:
# # # #                     try:
# # # #                         # Attempt decryption with an empty password
# # # #                         if pdf_reader.decrypt('') == 0: # 0 indicates failure
# # # #                              print(f"Warning: Could not decrypt PDF {file_path}")
# # # #                              return "", [] # Return empty if decryption fails
# # # #                     except Exception as decrypt_err:
# # # #                          print(f"Error decrypting PDF {file_path}: {decrypt_err}")
# # # #                          return "", [] # Return empty on decryption error
# # # #
# # # #                 text_content = ""
# # # #                 for page in pdf_reader.pages:
# # # #                     page_text = page.extract_text()
# # # #                     if page_text: # Check if text extraction was successful
# # # #                         text_content += page_text + "\n" # Add newline between pages
# # # #                 print(f"PyPDF2 extracted text length: {len(text_content)}")
# # # #             return text_content, []
# # # #         except Exception as e_pypdf:
# # # #             print(f"Error during fallback PDF text extraction with PyPDF2: {e_pypdf}")
# # # #             return "", []
# # # #
# # # #     doc = None # Initialize doc to None
# # # #     try:
# # # #         doc = fitz.open(file_path)
# # # #         for page_num in range(len(doc)):
# # # #             page = doc.load_page(page_num)
# # # #             text_content += page.get_text("text") + "\n" # Get plain text
# # # #             links = page.get_links()
# # # #             for link in links:
# # # #                 # Check if 'uri' key exists and kind is URI link
# # # #                 if link.get('kind') == fitz.LINK_URI and 'uri' in link:
# # # #                     uri = link['uri']
# # # #                     if uri and uri not in extracted_links:
# # # #                         # Basic validation for URI format
# # # #                         if isinstance(uri, str) and (uri.startswith('http') or uri.startswith('mailto:')):
# # # #                             extracted_links.append(uri)
# # # #                         elif isinstance(uri, str) and '@' in uri and '.' in uri and not uri.startswith('mailto:'):
# # # #                             # Attempt to fix potential email links missing mailto:
# # # #                              if re.match(r"[^@]+@[^@]+\.[^@]+", uri): # Basic email format check
# # # #                                 extracted_links.append(f"mailto:{uri}")
# # # #     except Exception as e_fitz:
# # # #         print(f"Error extracting text/links from PDF with PyMuPDF: {e_fitz}")
# # # #         # If PyMuPDF fails, potentially fallback to PyPDF2 again, or return empty
# # # #         return "", [] # Return empty on PyMuPDF error
# # # #     finally:
# # # #         if doc:
# # # #             try: doc.close()
# # # #             except Exception as e_close: print(f"Error closing PDF document: {e_close}")
# # # #
# # # #     print(f"PyMuPDF extracted text length: {len(text_content)}")
# # # #     print(f"Extracted Links: {extracted_links}")
# # # #     return text_content, extracted_links
# # # #
# # # # def extract_text_from_docx(file_path):
# # # #     text = "";
# # # #     try:
# # # #         doc = docx.Document(file_path)
# # # #         for para in doc.paragraphs: text += para.text + "\n"
# # # #     except Exception as e: print(f"Error extracting text from DOCX: {str(e)}")
# # # #     return text
# # # #
# # # # def extract_text_from_txt(file_path):
# # # #     text = "";
# # # #     try:
# # # #         # Try UTF-8 first, ignore errors
# # # #         with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
# # # #              text = file.read()
# # # #     except Exception as e:
# # # #         print(f"Error reading TXT file {file_path}: {str(e)}")
# # # #         # Optionally try other encodings here if needed, e.g., latin-1
# # # #         # try:
# # # #         #     with open(file_path, 'r', encoding='latin-1', errors='ignore') as file:
# # # #         #         text = file.read()
# # # #         # except Exception as e_latin:
# # # #         #     print(f"Error reading TXT with latin-1: {e_latin}")
# # # #     return text
# # # #
# # # # def extract_text_from_resume(file_obj):
# # # #     """Extract text and potentially links from various resume file formats."""
# # # #     if not file_obj or not hasattr(file_obj, 'name'):
# # # #         print("Error: Invalid file object passed to extract_text_from_resume.")
# # # #         # Return empty string and empty list for consistency
# # # #         return "", []
# # # #
# # # #     # Ensure file object is reset if it has been read before
# # # #     if hasattr(file_obj, 'seek'):
# # # #         try:
# # # #             file_obj.seek(0)
# # # #         except Exception as e:
# # # #             print(f"Warning: Could not seek file object - {e}")
# # # #
# # # #     file_path = None # Initialize file_path
# # # #     suffix = os.path.splitext(file_obj.name)[1].lower()
# # # #     # Use a 'with' statement for better file handling
# # # #     try:
# # # #         # Create temp file explicitly to manage deletion better
# # # #         fd, file_path = tempfile.mkstemp(suffix=suffix)
# # # #         os.close(fd) # Close the file descriptor
# # # #
# # # #         with open(file_path, 'wb') as temp_file_write:
# # # #              for chunk in file_obj.chunks():
# # # #                  temp_file_write.write(chunk)
# # # #
# # # #         text = ""; links = []
# # # #         file_ext = suffix
# # # #         if file_ext == '.pdf':
# # # #             text, links = extract_text_and_links_from_pdf(file_path) # Use new function
# # # #         elif file_ext in ['.docx', '.doc']:
# # # #             text = extract_text_from_docx(file_path)
# # # #         elif file_ext in ['.txt', '.odt']: # Consider adding specific ODT handling if needed
# # # #             text = extract_text_from_txt(file_path)
# # # #         else:
# # # #             print(f"Unsupported file format: {file_ext}")
# # # #         # Ensure text is always a string, even if extraction fails
# # # #         if not isinstance(text, str):
# # # #              text = ""
# # # #
# # # #     except Exception as e:
# # # #         print(f"Error during text extraction setup or file write: {e}")
# # # #         text = ""
# # # #         links = []
# # # #     finally:
# # # #         # Ensure temporary file is deleted
# # # #         if file_path and os.path.exists(file_path):
# # # #             try:
# # # #                 os.unlink(file_path)
# # # #             except Exception as e_unlink:
# # # #                  print(f"Error deleting temporary file {file_path}: {e_unlink}")
# # # #
# # # #     # Add a check for very short text which might indicate an extraction issue
# # # #     if len(text) < 50: # Adjusted threshold
# # # #          print(f"Warning: Extracted text is very short ({len(text)} chars). Extraction might have failed.")
# # # #
# # # #     return text, links # Return both text and links
# # # #
# # # #
# # # # # --- AI Parsing Function (Prompt Adjusted for Links) ---
# # # # def parse_resume_with_ai(resume_text, extracted_links=None, ai_engine='chatgpt'):
# # # #     """Parse resume text using AI to extract structured data, considering extracted links."""
# # # #     # Add a check for minimal text length before attempting AI parsing
# # # #     if not resume_text or len(resume_text.strip()) < 50: # Use same threshold as extraction
# # # #         print("Error: Resume text is too short or empty for AI parsing. Falling back to basic parsing.")
# # # #         return basic_resume_parsing(resume_text) # Return default dict
# # # #
# # # #     # Default to empty list if None
# # # #     if extracted_links is None:
# # # #         extracted_links = []
# # # #
# # # #     context_text = resume_text
# # # #     if extracted_links:
# # # #         context_text += "\n\n--- Extracted Links (Prioritize these for URLs/Email) ---\n"
# # # #         for link in extracted_links: context_text += f"- {link}\n"
# # # #         context_text += "--- End Extracted Links ---\n"
# # # #
# # # #     # --- REFINED PROMPT ---
# # # #     prompt = f"""
# # # #     Extract information from the following resume text and format it STRICTLY as a valid JSON object.
# # # #     Prioritize using URLs/Emails found in the 'Extracted Links' section if available and relevant, otherwise use URLs/Emails found directly in the text.
# # # #     Use the EXACT keys specified below. Use `null` for optional fields that are missing, DO NOT omit the key. Ensure all string values are enclosed in double quotes.
# # # #
# # # #     JSON Structure:
# # # #     {{
# # # #       "Personal Information": {{
# # # #         "First name": "string", "Middle name": "string | null", "Last name": "string",
# # # #         "Email": "string (use mailto: links or email addresses found)", "Phone number": "string",
# # # #         "Address": "string (e.g., 'City, State/Province') | null", "LinkedIn URL": "string (full URL) | null",
# # # #         "GitHub URL": "string (full URL) | null", "Portfolio URL": "string (full URL) | null"
# # # #       }},
# # # #       "Professional Summary": "string | null",
# # # #       "Skills": [ {{ "Skill name": "string", "Skill type": "string (technical, soft, language, tool)", "Estimated proficiency level": "integer (0-100)" }} ] | null,
# # # #       "Work Experience": [ {{ "Job title": "string", "Employer/Company name": "string", "Location": "string | null", "Start date": "string | null (YYYY-MM-DD)", "End date": "string | null (YYYY-MM-DD)", "Is current job": "boolean", "Bullet points": ["string", ...] | null }} ] | null,
# # # #       "Education": [ {{ "School name": "string", "Location": "string | null", "Degree": "string", "Degree type": "string (high_school, associate, bachelor, master, doctorate, other)", "Field of study": "string | null", "Graduation date": "string | null (YYYY-MM-DD)", "GPA": "float | null" }} ] | null,
# # # #       "Projects": [ {{ "Project name": "string", "Summary/description": "string | null", "Start date": "string | null (YYYY-MM-DD)", "Completion date": "string | null (YYYY-MM-DD)", "Project URL": "string | null", "GitHub URL": "string | null", "Bullet points": ["string", ...] | null }} ] | null,
# # # #       "Certifications": [ {{ "Name": "string", "Institute/Issuing organization": "string | null", "Completion date": "string | null (YYYY-MM-DD)", "Expiration date": "string | null (YYYY-MM-DD)", "Score": "string | null", "URL/Link": "string | null", "Description": "string | null" }} ] | null,
# # # #       "Languages": [ {{ "Language name": "string", "Proficiency": "string (basic, intermediate, advanced, native)" }} ] | null,
# # # #       "Additional sections": [ {{ "Section name": "string", "Completion date": "string | null (YYYY-MM-DD)", "Bullet points": ["string", ...] | null, "Description": "string | null", "URL/Link": "string | null", "Institution name": "string | null" }} ] | null
# # # #     }}
# # # #
# # # #     Ensure dates are YYYY-MM-DD format if possible, otherwise use text. Format Address/Location as "City, State/Province". Use key "Bullet points". Check Extracted Links section for relevant URLs/Emails. Output only valid JSON.
# # # #
# # # #     Resume Text & Links:
# # # #     ---
# # # #     {context_text}
# # # #     ---
# # # #     """
# # # #
# # # #     result = None
# # # #     try:
# # # #         # Clear previous errors if any (optional, depends on how errors are handled)
# # # #         # clear_previous_parsing_error()
# # # #
# # # #         if ai_engine == 'chatgpt' and openai.api_key:
# # # #             print("--- Using ChatGPT for parsing ---")
# # # #             response = openai.chat.completions.create(
# # # #                 model=getattr(settings, 'OPENAI_MODEL', 'gpt-4o-mini'), # Use a capable model
# # # #                 messages=[
# # # #                     {"role": "system", "content": "Parse resume text into the specified JSON format. Ensure all values are correct types (string, number, boolean, null, array). Prioritize URLs/Emails from 'Extracted Links' section. Output only valid JSON."},
# # # #                     {"role": "user", "content": prompt}
# # # #                 ],
# # # #                 temperature=0.1,
# # # #                 max_tokens=4000, # Increase max tokens if resumes are long
# # # #                 response_format={ "type": "json_object" } # Request JSON output format
# # # #             )
# # # #             result = response.choices[0].message.content
# # # #         elif ai_engine == 'gemini' and genai_api_key:
# # # #             print("--- Using Gemini for parsing ---")
# # # #             ai_settings = getattr(settings, 'AI_SETTINGS', {})
# # # #             gemini_config = ai_settings.get('gemini', {})
# # # #             model_name = gemini_config.get('model', 'gemini-1.5-flash') # Use Flash or Pro
# # # #             model = genai.GenerativeModel(model_name)
# # # #             response = model.generate_content(
# # # #                 prompt,
# # # #                 generation_config=genai.types.GenerationConfig(
# # # #                     temperature=0.2,
# # # #                     response_mime_type="application/json" # Request JSON output
# # # #                 )
# # # #             )
# # # #             result = response.text
# # # #         else:
# # # #             print("Warning: No valid AI engine configured/API key missing. Falling back to basic parsing.")
# # # #             return basic_resume_parsing(resume_text)
# # # #
# # # #         if not result:
# # # #             print("Error: AI returned an empty response. Falling back.")
# # # #             return basic_resume_parsing(resume_text)
# # # #
# # # #         # Clean the response (remove potential markdown/comments)
# # # #         result = re.sub(r'^```json\s*|\s*```$', '', result.strip(), flags=re.MULTILINE)
# # # #         result = re.sub(r'^\s*//.*$', '', result, flags=re.MULTILINE) # Remove potential comments
# # # #
# # # #         # Attempt to load the JSON
# # # #         parsed_data = json.loads(result)
# # # #         print("--- AI Parsing Successful ---")
# # # #         print(parsed_data)
# # # #         # **Crucial Check:** Ensure the top-level structure is a dictionary
# # # #         if not isinstance(parsed_data, dict):
# # # #             print(f"Error: Parsed data is not a dictionary after JSON load (type: {type(parsed_data)}). Falling back.")
# # # #             return basic_resume_parsing(resume_text)
# # # #
# # # #         return parsed_data
# # # #
# # # #     except json.JSONDecodeError as json_err:
# # # #         print(f"Error decoding JSON response from AI: {json_err}")
# # # #         print(f"AI Response causing error:\n>>>\n{result}\n<<<")
# # # #         # Maybe try to fix common JSON errors here or fallback
# # # #         return basic_resume_parsing(resume_text) # Fallback on JSON error
# # # #     except Exception as e:
# # # #         print(f"Error parsing resume with AI ({ai_engine}): {str(e)}")
# # # #         traceback.print_exc()
# # # #         return basic_resume_parsing(resume_text) # Fallback on other errors
# # # #
# # # #
# # # # # --- Basic Fallback Parser (Returns Dict) ---
# # # # def basic_resume_parsing(resume_text):
# # # #     """Basic fallback parsing function if AI parsing fails. Returns dictionary."""
# # # #     # Ensure resume_text is a string
# # # #     if not isinstance(resume_text, str):
# # # #         resume_text = str(resume_text) # Attempt to convert
# # # #
# # # #     email=None; phone=None; linkedin=None; github=None; portfolio=None; first_name=""; last_name=""; middle_name=None
# # # #     try:
# # # #         email_pattern=r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'; emails=re.findall(email_pattern, resume_text); email=emails[0] if emails else None
# # # #     except Exception as e: print(f"Error parsing email: {e}")
# # # #     try:
# # # #         phone_pattern=r'\b(?:\+\d{1,3}[-\.\s]?)?\(?\d{3}\)?[-\.\s]?\d{3}[-\.\s]?\d{4}\b'; phones=re.findall(phone_pattern, resume_text); phone=phones[0] if phones else None
# # # #     except Exception as e: print(f"Error parsing phone: {e}")
# # # #     try:
# # # #         url_pattern=r'https?://(?:www\.)?[-a-zA-Z0-9@:%._\+~#=]{1,256}\.[a-zA-Z0-9()]{1,6}\b(?:[-a-zA-Z0-9()@:%_\+.~#?&//=]*)'; urls=re.findall(url_pattern, resume_text)
# # # #         for url in urls:
# # # #             try:
# # # #                 parsed=urlparse(url); scheme=parsed.scheme or 'https'; netloc=parsed.netloc.replace('www.',''); path=parsed.path; clean_url=f"{scheme}://{netloc}{path}".rstrip('/')
# # # #                 if 'linkedin.com/in/' in clean_url and not linkedin:
# # # #                     linkedin=clean_url
# # # #                 elif 'github.com/' in clean_url and not github:
# # # #                     github=clean_url
# # # #                 elif not any(x in clean_url for x in ['linkedin.com','github.com']) and not portfolio:
# # # #                     # Basic check to avoid common non-portfolio URLs
# # # #                     if not any(ext in clean_url for ext in ['.pdf', '.png', '.jpg', '.jpeg', '.gif']):
# # # #                         portfolio=clean_url
# # # #             except ValueError:
# # # #                  print(f"Skipping malformed URL in basic_resume_parsing: {url}")
# # # #                  continue
# # # #     except Exception as e: print(f"Error parsing URLs: {e}")
# # # #     try:
# # # #         lines=resume_text.strip().split('\n'); name=lines[0].strip() if lines else ""
# # # #         name_parts=name.split(); first_name=name_parts[0] if name_parts else ""; last_name=name_parts[-1] if len(name_parts)>1 else ""; middle_name=" ".join(name_parts[1:-1]) if len(name_parts)>2 else None
# # # #     except Exception as e: print(f"Error parsing name: {e}")
# # # #
# # # #     # Return the default dictionary structure
# # # #     # Ensure all top-level keys are present even if empty
# # # #     return {
# # # #         "Personal Information": {
# # # #             "First name": first_name or None,
# # # #             "Middle name": middle_name,
# # # #             "Last name": last_name or None,
# # # #             "Email": email,
# # # #             "Phone number": phone,
# # # #             "Address": None,
# # # #             "LinkedIn URL": linkedin,
# # # #             "GitHub URL": github,
# # # #             "Portfolio URL": portfolio
# # # #         },
# # # #         "Professional Summary": None,
# # # #         "Skills": [], # Use empty list instead of None for consistency
# # # #         "Work Experience": [],
# # # #         "Education": [],
# # # #         "Projects": [],
# # # #         "Certifications": [],
# # # #         "Languages": [],
# # # #         "Additional sections": []
# # # #     }
# # # #
# # # #
# # # # # --- Helper Functions ---
# # # # def format_date(date_str):
# # # #     """Convert various date string formats to YYYY-MM-DD."""
# # # #     if not date_str or not isinstance(date_str, str) or str(date_str).lower() == 'present':
# # # #         return None
# # # #     date_str = str(date_str).strip()
# # # #     # Order formats from more specific to less specific
# # # #     formats = [
# # # #         '%Y-%m-%d', '%m/%d/%Y', '%m-%d-%Y', '%d-%b-%Y', '%d-%B-%Y',
# # # #         '%Y/%m/%d', '%d %b %Y', '%d %B %Y', '%b %d, %Y', '%B %d, %Y',
# # # #         '%Y-%m', '%m/%Y', '%b %Y', '%B %Y', '%Y',
# # # #     ]
# # # #     for fmt in formats:
# # # #         try:
# # # #             dt = datetime.strptime(date_str, fmt)
# # # #             # Handle formats that only provide year and month or just year
# # # #             if fmt in ['%Y-%m', '%m/%Y', '%b %Y', '%B %Y']:
# # # #                 return dt.strftime('%Y-%m-01') # Default to the 1st of the month
# # # #             elif fmt == '%Y':
# # # #                 return dt.strftime('%Y-01-01') # Default to Jan 1st of the year
# # # #             else:
# # # #                 return dt.strftime('%Y-%m-%d') # Full date
# # # #         except ValueError:
# # # #             continue # Try the next format
# # # #
# # # #     # Fallback attempts (keep simple for now)
# # # #     try:
# # # #         match_month_year = re.match(r'([A-Za-z]{3,})\s+(\d{4})', date_str)
# # # #         if match_month_year:
# # # #             month_str, year_str = match_month_year.groups()
# # # #             # Try common month abbreviations/names
# # # #             try:
# # # #                 dt = datetime.strptime(f"{month_str} 1 {year_str}", '%b 1 %Y')
# # # #                 return dt.strftime('%Y-%m-01')
# # # #             except ValueError:
# # # #                 try:
# # # #                     dt = datetime.strptime(f"{month_str} 1 {year_str}", '%B 1 %Y')
# # # #                     return dt.strftime('%Y-%m-01')
# # # #                 except ValueError: pass # Give up on month/year text
# # # #
# # # #         # Attempt to parse just YYYY as final fallback
# # # #         match_year = re.match(r'(\d{4})', date_str)
# # # #         if match_year:
# # # #             year = match_year.group(1)
# # # #             return datetime(int(year), 1, 1).strftime('%Y-01-01')
# # # #
# # # #     except Exception as e:
# # # #          print(f"Error during fallback date parsing for '{date_str}': {e}")
# # # #
# # # #     print(f"Warning: Could not parse date string: {date_str}")
# # # #     return None # Return None if no format matches
# # # #
# # # # def format_url(url_str):
# # # #     """Adds https:// if missing and cleans the URL."""
# # # #     if not url_str or not isinstance(url_str, str): return None; url_str = url_str.strip();
# # # #     if not url_str: return None;
# # # #     try:
# # # #         # Add scheme if missing, default to https
# # # #         if not re.match(r"^(?:[a-z]+:)?//", url_str, re.IGNORECASE) and '.' in url_str:
# # # #              url_str = 'https://' + url_str
# # # #
# # # #         parsed = urlparse(url_str);
# # # #         scheme = parsed.scheme; netloc = parsed.netloc; path = parsed.path or ''; # Ensure path is a string
# # # #
# # # #         # Basic validation
# # # #         if not scheme or not netloc or '.' not in netloc:
# # # #             # If it looks like an email, add mailto:
# # # #             if '@' in url_str and '.' in url_str and not url_str.startswith('mailto:'):
# # # #                 if re.match(r"[^@]+@[^@]+\.[^@]+", url_str): # Basic email format check
# # # #                     return f"mailto:{url_str}"
# # # #             return None # Invalid URL structure
# # # #
# # # #         # Reconstruct URL, remove trailing slash if path is more than just '/'
# # # #         clean_url = f"{scheme}://{netloc}{path}"
# # # #         if clean_url.endswith('/') and len(path) > 1:
# # # #             clean_url = clean_url.rstrip('/')
# # # #         return clean_url
# # # #     except Exception as e: print(f"Error formatting URL '{url_str}': {e}"); return None
# # # #
# # # # def format_location(location_data):
# # # #     """Formats location data (string, dict, list, tuple) into a string."""
# # # #     if isinstance(location_data, dict):
# # # #         city = location_data.get('City', '')
# # # #         state = location_data.get('State/Province', '')
# # # #         country = location_data.get('Country', '')
# # # #         # Filter out None or empty strings before joining
# # # #         parts = [str(part).strip() for part in [city, state, country] if part and isinstance(part, str) and str(part).strip()]
# # # #         return ", ".join(parts) if parts else ""
# # # #     elif isinstance(location_data, str):
# # # #         return location_data.strip()
# # # #     elif isinstance(location_data, (list, tuple)): # Handle list/tuple case
# # # #         # Join elements if they are strings, filter out non-strings/empty
# # # #         parts = [str(item).strip() for item in location_data if item and isinstance(item, str) and str(item).strip()]
# # # #         return ", ".join(parts) if parts else ""
# # # #     else:
# # # #         # Return empty string for other types (like None, int, float, etc.)
# # # #         return ""
# # # #
# # # # def safe_strip(value, default=''):
# # # #     """Safely strips a value if it's a string, joins if list/tuple, otherwise returns default."""
# # # #     if isinstance(value, str):
# # # #         return value.strip()
# # # #     elif isinstance(value, (list, tuple)):
# # # #         # Join elements with a space if it's a list/tuple of strings
# # # #         # Filter out non-string items before joining
# # # #         string_items = [str(item).strip() for item in value if isinstance(item, str) and str(item).strip()]
# # # #         return " ".join(string_items) if string_items else default
# # # #     # Handle cases where value might be None, int, float, etc.
# # # #     elif value is None:
# # # #         return default
# # # #     else:
# # # #         # Attempt to convert other types to string, then return default if fails
# # # #         try:
# # # #             return str(value).strip()
# # # #         except:
# # # #             return default
# # # #
# # # # # --- create_resume_from_parsed_data (Obsolete - Remove if unused or implement) ---
# # # # # def create_resume_from_parsed_data(user, parsed_data, template_id=None):
# # # # #     # This function seems to be defined elsewhere or needs implementation
# # # # #     pass